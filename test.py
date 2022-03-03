from datetime import datetime

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.text import font

from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)


def isDate(dateString: str) -> bool:
    try:
        dateParsed = datetime.strptime(dateString, "%d/%m/%y")
        return True
    except Exception as e:
        return False


def stringBeginsWithDate(string: str) -> bool:
    dateString = string[:8]

    return isDate(dateString = dateString)


def isLorenaMessage(string: str) -> bool:
    header = string[:40]

    return 'maría lorena' in header


def getHeader(string: str) -> str:
    colonIndexes = [pos for pos, char in enumerate(string) if char == ':']
    try:
        indexesOfLastHeaderColon = colonIndexes[1]

        header = string[0:indexesOfLastHeaderColon + 1]
        return header

    except:
        indexesOfLastHeaderColon = colonIndexes[0]
        header = string[0:indexesOfLastHeaderColon + 1]
        return header


def getBody(string: str) -> str:
    colonIndexes = [pos for pos, char in enumerate(string) if char == ':']
    try:
        indexesOfLastHeaderColon = colonIndexes[1]

        header = string[indexesOfLastHeaderColon + 1:]
        return header
    except:
        indexesOfLastHeaderColon = colonIndexes[0]
        header = string[indexesOfLastHeaderColon + 1:]
        return header


if __name__ == '__main__':

    originalDoc = Document('chat.docx')
    editedDoc = Document()

    dates = []

    p = editedDoc.add_heading('Chat Whatsapp with Andrea Varona', 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    isFirstDate = True

    isLorenaMsg = False

    for i in originalDoc.paragraphs:

        if stringBeginsWithDate(i.text):
            isLorenaMsg = isLorenaMessage(i.text)

            if isLorenaMsg:
                highlightColor = WD_COLOR_INDEX.YELLOW
            else:
                highlightColor = WD_COLOR_INDEX.GRAY_25

            p = editedDoc.add_paragraph()
            insertHR(p)
            date = i.text[:8]
            if date not in dates:
                dates.append(date)
                if not isFirstDate:
                    editedDoc.add_page_break()
                isFirstDate = False
                h = editedDoc.add_heading(date, level = 1)
                h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            font.bold = True
            p = editedDoc.add_paragraph()
            run = p.add_run(getHeader(i.text))
            run.bold = True
            run.font.highlight_color = highlightColor
            run = p.add_run(getBody(i.text))
            run.font.highlight_color = highlightColor

        else:
            p.add_run().add_break()
            run = p.add_run(i.text)
            run.font.highlight_color = highlightColor

    editedDoc.save('test.docx')
